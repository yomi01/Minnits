"""
Output Formatter Module

This module provides functionality to save summaries in multiple formats:
- JSON (original format)
- DOCX (Microsoft Word document)
- TXT (plain text)
"""

import os
import json
from typing import Dict, Any, Optional
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class SummaryFormatter:
    """
    A class to format and save summaries in different output formats.
    """
    
    def __init__(self):
        """Initialize the summary formatter."""
        self.supported_formats = ['json', 'txt']
        if DOCX_AVAILABLE:
            self.supported_formats.append('docx')
    
    def get_supported_formats(self) -> list:
        """Return list of supported output formats."""
        return self.supported_formats.copy()
    
    def save_summary(self, summary_data: Dict[str, Any], output_path: str, 
                    format_type: str = 'json', include_metadata: bool = True) -> str:
        """
        Save summary in the specified format.
        
        Args:
            summary_data: Dictionary containing summary information
            output_path: Base path for output file (without extension)
            format_type: Output format ('json', 'docx', 'txt')
            include_metadata: Whether to include metadata in output
            
        Returns:
            Full path to the saved file
        """
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}. Supported: {self.supported_formats}")
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Generate filename with proper extension
        base_name = os.path.splitext(output_path)[0]
        file_path = f"{base_name}.{format_type}"
        
        # Save in the requested format
        if format_type == 'json':
            return self._save_json(summary_data, file_path, include_metadata)
        elif format_type == 'docx':
            return self._save_docx(summary_data, file_path, include_metadata)
        elif format_type == 'txt':
            return self._save_txt(summary_data, file_path, include_metadata)
        else:
            raise ValueError(f"Handler not implemented for format: {format_type}")
    
    def save_multiple_formats(self, summary_data: Dict[str, Any], output_path: str,
                             formats: list = None, include_metadata: bool = True) -> Dict[str, str]:
        """
        Save summary in multiple formats.
        
        Args:
            summary_data: Dictionary containing summary information
            output_path: Base path for output files (without extension)
            formats: List of formats to save (defaults to all supported)
            include_metadata: Whether to include metadata in output
            
        Returns:
            Dictionary mapping format to file path
        """
        if formats is None:
            formats = self.supported_formats
        
        saved_files = {}
        for format_type in formats:
            if format_type in self.supported_formats:
                try:
                    file_path = self.save_summary(summary_data, output_path, format_type, include_metadata)
                    saved_files[format_type] = file_path
                except Exception as e:
                    print(f"Warning: Failed to save {format_type} format: {e}")
        
        return saved_files
    
    def _save_json(self, summary_data: Dict[str, Any], file_path: str, include_metadata: bool) -> str:
        """Save summary as JSON file."""
        output_data = summary_data.copy()
        
        if include_metadata:
            output_data["export_metadata"] = {
                "exported_at": datetime.now().isoformat(),
                "format": "json",
                "version": "1.0"
            }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        
        return file_path
    
    def _save_docx(self, summary_data: Dict[str, Any], file_path: str, include_metadata: bool) -> str:
        """Save summary as DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required to save DOCX files. Install it with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Summary Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Source information
        if 'source' in summary_data:
            doc.add_heading('Source Document', level=1)
            source_para = doc.add_paragraph()
            source_para.add_run('Document: ').bold = True
            source_para.add_run(str(summary_data['source']))
        
        if 'processed_at' in summary_data:
            processed_para = doc.add_paragraph()
            processed_para.add_run('Processed: ').bold = True
            processed_para.add_run(str(summary_data['processed_at']))
        
        # Check if we have detailed format
        has_detailed_format = summary_data.get("format_used") == "detailed"
        
        if has_detailed_format:
            # Use detailed meeting minutes format
            if summary_data.get("meeting_overview"):
                doc.add_heading('Meeting Overview', level=1)
                doc.add_paragraph(summary_data["meeting_overview"])
            
            if summary_data.get("discussion_highlights"):
                doc.add_heading('Discussion Highlights', level=1)
                doc.add_paragraph(summary_data["discussion_highlights"])
            
            if summary_data.get("decisions_made"):
                doc.add_heading('Decisions Made', level=1)
                decisions = summary_data["decisions_made"]
                if isinstance(decisions, list):
                    for decision in decisions:
                        para = doc.add_paragraph(style='List Bullet')
                        para.text = str(decision)
                else:
                    doc.add_paragraph(str(decisions))
            
            if summary_data.get("action_points"):
                doc.add_heading('Action Items', level=1)
                action_points = summary_data["action_points"]
                if isinstance(action_points, list):
                    for action in action_points:
                        para = doc.add_paragraph(style='List Bullet')
                        para.text = str(action)
                else:
                    doc.add_paragraph(str(action_points))
            
            if summary_data.get("next_steps"):
                doc.add_heading('Next Steps/Follow-Up', level=1)
                doc.add_paragraph(summary_data["next_steps"])
        else:
            # Use simple format (backward compatibility)
            doc.add_heading('Summary', level=1)
            summary_text = summary_data.get('summary', 'No summary available')
            doc.add_paragraph(summary_text)
            
            # Action points section
            if 'action_points' in summary_data and summary_data['action_points']:
                doc.add_heading('Action Points', level=1)
                action_points = summary_data['action_points']
                
                if isinstance(action_points, list):
                    for i, action in enumerate(action_points, 1):
                        para = doc.add_paragraph(style='List Number')
                        para.text = str(action)
                else:
                    doc.add_paragraph(str(action_points))
        
        # Key points section
        if 'key_points' in summary_data and summary_data['key_points']:
            doc.add_heading('Key Points', level=1)
            key_points = summary_data['key_points']
            
            if isinstance(key_points, list):
                for point in key_points:
                    para = doc.add_paragraph(style='List Bullet')
                    para.text = str(point)
            else:
                doc.add_paragraph(str(key_points))
        
        # Participants section
        if 'participants' in summary_data and summary_data['participants']:
            doc.add_heading('Participants', level=1)
            participants = summary_data['participants']
            
            if isinstance(participants, list):
                for participant in participants:
                    para = doc.add_paragraph(style='List Bullet')
                    para.text = str(participant)
            else:
                doc.add_paragraph(str(participants))
        
        # Statistics section
        if 'metadata' in summary_data:
            doc.add_heading('Document Statistics', level=1)
            meta = summary_data['metadata']
            
            stats_para = doc.add_paragraph()
            if 'word_count' in meta:
                stats_para.add_run(f"Word Count: {meta['word_count']:,}\n")
            if 'character_count' in meta:
                stats_para.add_run(f"Character Count: {meta['character_count']:,}\n")
            if 'total_paragraphs' in meta:
                stats_para.add_run(f"Paragraphs: {meta['total_paragraphs']}\n")
        
        # Export metadata
        if include_metadata:
            doc.add_page_break()
            doc.add_heading('Export Information', level=1)
            meta_para = doc.add_paragraph()
            meta_para.add_run('Exported: ').bold = True
            meta_para.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
            meta_para.add_run('\nFormat: ').bold = True
            meta_para.add_run('Microsoft Word Document (.docx)')
            meta_para.add_run('\nGenerated by: ').bold = True
            meta_para.add_run('Note Summarizer App')
        
        doc.save(file_path)
        return file_path
    
    def _save_txt(self, summary_data: Dict[str, Any], file_path: str, include_metadata: bool) -> str:
        """Save summary as plain text file."""
        lines = []
        
        # Title
        lines.append("=" * 60)
        lines.append("SUMMARY REPORT")
        lines.append("=" * 60)
        lines.append("")
        
        # Source information
        if 'source' in summary_data:
            lines.append(f"Source Document: {summary_data['source']}")
        if 'processed_at' in summary_data:
            lines.append(f"Processed: {summary_data['processed_at']}")
        lines.append("")
        
        # Check if we have detailed format
        has_detailed_format = summary_data.get("format_used") == "detailed"
        
        if has_detailed_format:
            # Use detailed meeting minutes format
            if summary_data.get("meeting_overview"):
                lines.append("MEETING OVERVIEW")
                lines.append("-" * 40)
                lines.append(summary_data["meeting_overview"])
                lines.append("")
            
            if summary_data.get("discussion_highlights"):
                lines.append("DISCUSSION HIGHLIGHTS")
                lines.append("-" * 40)
                lines.append(summary_data["discussion_highlights"])
                lines.append("")
            
            if summary_data.get("decisions_made"):
                lines.append("DECISIONS MADE")
                lines.append("-" * 40)
                decisions = summary_data["decisions_made"]
                if isinstance(decisions, list):
                    for decision in decisions:
                        lines.append(f"• {decision}")
                else:
                    lines.append(str(decisions))
                lines.append("")
            
            if summary_data.get("action_points"):
                lines.append("ACTION ITEMS")
                lines.append("-" * 40)
                action_points = summary_data["action_points"]
                if isinstance(action_points, list):
                    for action in action_points:
                        lines.append(f"• {action}")
                else:
                    lines.append(str(action_points))
                lines.append("")
            
            if summary_data.get("next_steps"):
                lines.append("NEXT STEPS/FOLLOW-UP")
                lines.append("-" * 40)
                lines.append(summary_data["next_steps"])
                lines.append("")
        else:
            # Use simple format (backward compatibility)
            lines.append("SUMMARY")
            lines.append("-" * 40)
            summary_text = summary_data.get('summary', 'No summary available')
            lines.append(summary_text)
            lines.append("")
            
            # Action points section
            if 'action_points' in summary_data and summary_data['action_points']:
                lines.append("ACTION POINTS")
                lines.append("-" * 40)
                action_points = summary_data['action_points']
                
                if isinstance(action_points, list):
                    for i, action in enumerate(action_points, 1):
                        lines.append(f"{i}. {action}")
                else:
                    lines.append(str(action_points))
                lines.append("")
        
        # Key points section
        if 'key_points' in summary_data and summary_data['key_points']:
            lines.append("KEY POINTS")
            lines.append("-" * 40)
            key_points = summary_data['key_points']
            
            if isinstance(key_points, list):
                for point in key_points:
                    lines.append(f"• {point}")
            else:
                lines.append(str(key_points))
            lines.append("")
        
        # Participants section
        if 'participants' in summary_data and summary_data['participants']:
            lines.append("PARTICIPANTS")
            lines.append("-" * 40)
            participants = summary_data['participants']
            
            if isinstance(participants, list):
                for participant in participants:
                    lines.append(f"• {participant}")
            else:
                lines.append(str(participants))
            lines.append("")
        
        # Statistics section
        if 'metadata' in summary_data:
            lines.append("DOCUMENT STATISTICS")
            lines.append("-" * 40)
            meta = summary_data['metadata']
            
            if 'word_count' in meta:
                lines.append(f"Word Count: {meta['word_count']:,}")
            if 'character_count' in meta:
                lines.append(f"Character Count: {meta['character_count']:,}")
            if 'total_paragraphs' in meta:
                lines.append(f"Paragraphs: {meta['total_paragraphs']}")
            lines.append("")
        
        # Export metadata
        if include_metadata:
            lines.append("=" * 60)
            lines.append("EXPORT INFORMATION")
            lines.append("=" * 60)
            lines.append(f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            lines.append("Format: Plain Text (.txt)")
            lines.append("Generated by: Note Summarizer App")
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(lines))
        
        return file_path


def create_summary_formatter() -> SummaryFormatter:
    """Factory function to create a SummaryFormatter instance."""
    return SummaryFormatter()
